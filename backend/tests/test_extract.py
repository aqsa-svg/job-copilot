import io

import pytest

from app.extract import ExtractError, MAX_BYTES, extract_text


def test_txt():
    assert "Hello world" in extract_text("resume.txt", "text/plain", b"Hello world")


def test_docx():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("ML Engineer at Acme")
    buf = io.BytesIO()
    doc.save(buf)
    text = extract_text("resume.docx", None, buf.getvalue())
    assert "Jane Doe" in text and "Acme" in text


def test_docx_extracts_embedded_hyperlink():
    """A link 'wrapped' behind anchor text should still be recovered."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("See my profile")  # visible text has no URL
    doc.part.relate_to(
        "https://linkedin.com/in/jane",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    buf = io.BytesIO()
    doc.save(buf)
    text = extract_text("resume.docx", None, buf.getvalue())
    assert "linkedin.com/in/jane" in text
    assert "LINKS FOUND IN DOCUMENT" in text


def test_plain_text_urls_are_collected():
    text = extract_text("r.txt", "text/plain", b"Contact: https://github.com/jane and email")
    assert "https://github.com/jane" in text


def test_unsupported_binary_raises():
    with pytest.raises(ExtractError):
        extract_text("photo.jpg", "image/jpeg", b"\xff\xd8\xff\xe0\xff\xfe")


def test_legacy_doc_raises():
    with pytest.raises(ExtractError):
        extract_text("resume.doc", None, b"anything")


def test_too_large_raises():
    with pytest.raises(ExtractError):
        extract_text("resume.txt", "text/plain", b"x" * (MAX_BYTES + 1))


# ---- API upload endpoint (LLM mocked) ----
def test_parse_file_txt(client, monkeypatch):
    monkeypatch.setattr(
        "app.llm.parse_resume",
        lambda text: {"full_name": "From File", "seen": text[:8]},
    )
    r = client.post(
        "/api/profile/parse-file",
        files={"file": ("resume.txt", b"John Doe\nEngineer", "text/plain")},
    )
    assert r.status_code == 200
    assert r.json()["full_name"] == "From File"
    assert r.json()["seen"] == "John Doe"


def test_parse_file_empty(client):
    r = client.post(
        "/api/profile/parse-file",
        files={"file": ("resume.txt", b"", "text/plain")},
    )
    assert r.status_code == 400
