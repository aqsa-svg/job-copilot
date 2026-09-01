"""Extract plain text from an uploaded resume file (PDF / DOCX / TXT).

Besides visible text, this also recovers hyperlink URLs that are "wrapped"
behind anchor text (e.g. the word "LinkedIn" linked to a URL) — PDF link
annotations and DOCX relationships — plus any plain URLs in the text, and
appends them so the parser can populate the profile's links.
"""
from __future__ import annotations

import io
import re

MAX_BYTES = 8 * 1024 * 1024  # 8 MB cap

_URL_RE = re.compile(r"https?://[^\s)>\]}\"']+", re.IGNORECASE)


class ExtractError(Exception):
    """Raised when a file can't be read into usable text."""


def extract_text(filename: str, content_type: str | None, data: bytes) -> str:
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if len(data) > MAX_BYTES:
        raise ExtractError("File is too large (max 8 MB).")

    if name.endswith(".pdf") or "pdf" in ctype:
        text, links = _from_pdf(data)
    elif name.endswith(".docx") or "wordprocessingml" in ctype:
        text, links = _from_docx(data)
    elif name.endswith((".txt", ".md")) or ctype.startswith("text/"):
        text, links = data.decode("utf-8", errors="ignore"), []
    elif name.endswith(".doc"):
        raise ExtractError(
            "Legacy .doc isn't supported - save as .docx or PDF, or paste the text."
        )
    else:
        text = data.decode("utf-8", errors="ignore").strip()
        if not text:
            raise ExtractError("Unsupported file type. Upload a PDF, DOCX, or TXT.")
        links = []

    return _append_links(text, links)


def _append_links(text: str, embedded: list[str]) -> str:
    """Merge embedded + in-text URLs, dedupe, and append a LINKS section."""
    found = list(embedded) + _URL_RE.findall(text or "")
    seen: set[str] = set()
    uniq: list[str] = []
    for raw in found:
        u = raw.strip().rstrip(".,);]>'\"")
        if u and u.lower() not in seen:
            seen.add(u.lower())
            uniq.append(u)
    if uniq:
        text = f"{text or ''}\n\nLINKS FOUND IN DOCUMENT:\n" + "\n".join(uniq)
    return text


def _from_pdf(data: bytes) -> tuple[str, list[str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ExtractError("pypdf is not installed. Run: pip install -r requirements.txt") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # corrupt/encrypted PDF, etc.
        raise ExtractError(f"Could not read the PDF: {exc}") from exc

    parts: list[str] = []
    links: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            pass
        links.extend(_pdf_page_links(page))
    return "\n".join(parts).strip(), links


def _pdf_page_links(page) -> list[str]:
    """Pull /URI targets from a page's link annotations."""
    out: list[str] = []
    try:
        annots = page.get("/Annots")
    except Exception:
        annots = None
    if not annots:
        return out
    try:
        annots = list(annots)
    except TypeError:
        return out
    for a in annots:
        try:
            obj = a.get_object()
            action = obj.get("/A")
            if action is not None:
                uri = action.get_object().get("/URI")
                if uri:
                    out.append(str(uri))
        except Exception:
            continue
    return out


def _from_docx(data: bytes) -> tuple[str, list[str]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ExtractError("python-docx is not installed. Run: pip install -r requirements.txt") from exc
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractError(f"Could not read the .docx: {exc}") from exc

    lines = [p.text for p in doc.paragraphs]
    # Include table cell text too (resumes often use tables for layout).
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    # Embedded hyperlink targets live in the document relationships.
    links: list[str] = []
    try:
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype and getattr(rel, "is_external", False):
                links.append(rel.target_ref)
    except Exception:
        pass

    return "\n".join(lines).strip(), links
